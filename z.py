from docx import Document

# Good, now let's map M.docx's content into the same format as ETHICA.docx ("No | Indicators | Answers").

# Indicators list as in ETHICA doc
indicators = [
    "Paper Title",
    "Name of author(s)",
    "University of the authors",
    "The name of the journal used",
    "Year of publication",
    "The aim of the study",
    "The objective of the study/paper",
    "The research question(s) posed",
    "The hypotheses stated",
    "Type of moderator variable(s)",
    "Type of mediator variable(s)",
    "The theory(ies) employed",
    "Country used for the study",
    "Sectors/industry used for the study",
    "Research population",
    "Sample size",
    "Sampling method",
    "Qualitative/Quantitative/Mixed method",
    "Type of qualitative method",
    "Type of quantitative method",
    "Data collection method",
    "Findings – list them one by one",
    "Avenue for future studies – list them one by one",
    "Limitations and weaknesses of the study"
]

# Answers extracted from Paper Summary M.docx
answers = [
    "Ethical leadership and performance: The effect of follower individualism-collectivism",
    "Danna Booyens Strydom",
    "University of Pretoria, South Africa",
    "International Journal of Cross Cultural Management",
    "2021",
    "To explore how cultural value orientations—specifically individualism and collectivism—shape the relationship between ethical leadership and employee outcomes.",
    "1. To test whether ethical leadership improves employee performance and organizational citizenship behaviour (OCB).\n2. To determine whether horizontal/vertical forms of individualism and collectivism change (moderate) these effects.",
    "1. Does ethical leadership enhance employee performance and OCB?\n2. How do different cultural orientations of employees influence this relationship?",
    "H1: Ethical leadership positively relates to OCB.\nH2: Ethical leadership positively relates to task performance.\nH3: Horizontal individualism weakens the EL–OCB relationship.\nH4: Horizontal collectivism strengthens the EL–OCB relationship.\nH5: Vertical individualism weakens the EL–OCB relationship.\nH6: Vertical collectivism weakens the EL–OCB relationship.\nH7: Horizontal individualism weakens the EL–performance relationship.\nH8: Horizontal collectivism strengthens the EL–performance relationship.",
    "Cultural orientations: horizontal collectivism, vertical collectivism, horizontal individualism, vertical individualism.",
    "None directly tested (the focus was only on moderation).",
    "Social Cognitive Theory (Bandura)\nSocial Learning Theory (Bandura)\nSocial Exchange Theory (Blau)",
    "South Africa (study done in a multinational firm with branches across several African countries).",
    "Leisure service industry.",
    "Middle managers of the multinational company and their employees.",
    "352 respondents (after cleaning the data).",
    "Survey of all employees reporting to the selected managers (census-style, not random).",
    "Quantitative method",
    "Not applicable.",
    "Structural Equation Modelling (AMOS).",
    "1. Employee surveys (to measure ethical leadership and cultural orientations).\n2. Supervisor ratings (to assess OCB).\n3. Company records (objective performance scores).",
    "1. Ethical leadership is positively linked to both OCB and task performance.\n2. Horizontal individualism weakens the EL–OCB relationship.\n3. Horizontal collectivism strengthens the EL–OCB relationship.\n4. Vertical collectivism weakens the EL–OCB relationship.\n5. Vertical individualism showed no effect.\n6. Cultural orientation did not significantly change the EL–performance relationship.",
    "1. Use longitudinal or experimental designs to confirm causality.\n2. Conduct multi-country comparisons to explore differences across nations.\n3. Examine additional cultural values, such as power distance.\n4. Test in other industries beyond leisure services.\n5. Study leader cultural orientation and employee maturity.\n6. Explore ethical leadership’s effect on outcomes such as ethical climate.",
    "1. Cross-sectional design prevents strong causal claims.\n2. Sample came from only one company and industry.\n3. Only one cultural dimension (individualism–collectivism) was studied.\n4. Country-level comparisons were not possible due to limited respondents per country."
]

# Now build a new doc with the same style as ETHICA
new_doc = Document()

# Add student info placeholder like ETHICA doc (not given in M.docx, so we'll leave blank or generic)
new_doc.add_paragraph("Updated Summary - Formatted like ETHICA")

# Create table with 3 columns
table = new_doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No'
hdr_cells[1].text = 'Indicators'
hdr_cells[2].text = 'Answers'

# Fill rows
for i, (ind, ans) in enumerate(zip(indicators, answers), start=1):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = ind
    row_cells[2].text = ans

# Save
output_path = r"C:\Users\HP\Downloads\Paper_Summary_M_Formatted.docx"
new_doc.save(output_path)

output_path
